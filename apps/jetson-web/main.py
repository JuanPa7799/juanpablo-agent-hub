import hashlib
import hmac
import html
import ipaddress
import json
import os
import re
import secrets
import shutil
import socket
import sqlite3
import time
import uuid
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from urllib.parse import urljoin, urlparse

import httpx
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from fastapi import (
    Cookie,
    Depends,
    FastAPI,
    File,
    Form,
    Header,
    HTTPException,
    Response,
    UploadFile,
)
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


BASE_DIR = Path(__file__).resolve().parent
PUBLIC_DIR = BASE_DIR / "public"
DATA_DIR = BASE_DIR / "data"
DB_PATH = DATA_DIR / "agentes_datos.db"
EMPLOYMENT_DIR = DATA_DIR / "empleo"
DOCUMENTS_DIR = EMPLOYMENT_DIR / "documents"
ARTIFACTS_DIR = EMPLOYMENT_DIR / "artifacts"
ENV_PATH = BASE_DIR / ".env"
SESSION_COOKIE = "jetson_web_session"
MAX_UPLOAD_BYTES = 10 * 1024 * 1024
MAX_URL_BYTES = 2 * 1024 * 1024
ALLOWED_DOCUMENT_EXTENSIONS = {".docx", ".pdf", ".png", ".jpg", ".jpeg"}
ALLOWED_DOCUMENT_MIME = {
    ".docx": {"application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/octet-stream"},
    ".pdf": {"application/pdf", "application/octet-stream"},
    ".png": {"image/png", "application/octet-stream"},
    ".jpg": {"image/jpeg", "application/octet-stream"},
    ".jpeg": {"image/jpeg", "application/octet-stream"},
}


def load_env(path: Path) -> None:
    if not path.exists():
        return
    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        os.environ.setdefault(key.strip(), value.strip().strip("'").strip('"'))


load_env(ENV_PATH)

OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY", "")
OPENROUTER_MODEL = os.getenv("OPENROUTER_MODEL", "poolside/laguna-m.1:free")
WEB_USER = os.getenv("JETSON_WEB_USER", "juanpablo")
WEB_PASSWORD = os.getenv("JETSON_WEB_PASSWORD", "")
SESSION_SECRET = os.getenv("JETSON_WEB_SESSION_SECRET", "")
APP_TOKEN = os.getenv("JETSON_WEB_APP_TOKEN", "")
AUTH_ENABLED = os.getenv("JETSON_WEB_AUTH_ENABLED", "true").lower() != "false"
COOKIE_SECURE = os.getenv("JETSON_WEB_COOKIE_SECURE", "true").lower() != "false"


class LoginRequest(BaseModel):
    username: str
    password: str


class AIRequest(BaseModel):
    prompt: str
    system_prompt: Optional[str] = None
    is_json: bool = False
    model: Optional[str] = None



class OrchestratorRequest(BaseModel):
    prompt: str
    app_id: str
    session_id: Optional[str] = None
    agent_id: Optional[str] = None

class StateRequest(BaseModel):
    state: Dict[str, Any]


class BitacoraRequest(BaseModel):
    fecha: str
    contenido: str
    app_id: str = "predoctorado"


class DocumentPatch(BaseModel):
    category: Optional[str] = None
    status: Optional[str] = None


class ProfileRequest(BaseModel):
    profile: Dict[str, Any]


class VacancyParseRequest(BaseModel):
    url: str = ""
    text: str = ""
    profile_track: str = "general"
    modality: str = "remoto_hibrido"
    location: str = "Mexico / LATAM"
    language: str = "auto"
    level: str = "junior_mid"


class VacancyPatch(BaseModel):
    fields: Dict[str, Any]


class CVGenerateRequest(BaseModel):
    vacancy: Dict[str, Any]
    profile_track: str = "auto"
    language: str = "auto"


class GenerationPatch(BaseModel):
    content: Dict[str, Any]


app = FastAPI(title="Jetson Web API", version="2.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:8000",
        "http://127.0.0.1:8000",
        "http://localhost:8080",
        "http://127.0.0.1:8080",
        "http://192.168.1.230:8000",
        "https://juanpa7799.github.io",
        "https://JuanPa7799.github.io",
        "http://juanpablogc.com",
        "https://juanpablogc.com",
        "http://www.juanpablogc.com",
        "https://www.juanpablogc.com",
    ],
    allow_origin_regex=r"https://[-a-zA-Z0-9]+\.trycloudflare\.com",
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def now() -> int:
    return int(time.time())


def db() -> sqlite3.Connection:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(DB_PATH, timeout=20)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA synchronous=NORMAL")
    conn.execute("PRAGMA foreign_keys=ON")
    return conn


def init_db() -> None:
    DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)
    ARTIFACTS_DIR.mkdir(parents=True, exist_ok=True)
    with db() as conn:
        conn.executescript(
            """
            CREATE TABLE IF NOT EXISTS app_state (
                app_id TEXT PRIMARY KEY,
                state_json TEXT NOT NULL,
                updated_at INTEGER NOT NULL
            );
            CREATE TABLE IF NOT EXISTS bitacoras (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                app_id TEXT NOT NULL,
                fecha TEXT NOT NULL,
                contenido TEXT NOT NULL,
                created_at INTEGER NOT NULL
            );
            CREATE TABLE IF NOT EXISTS empleo_documents (
                id TEXT PRIMARY KEY,
                original_name TEXT NOT NULL,
                stored_name TEXT NOT NULL UNIQUE,
                extension TEXT NOT NULL,
                mime_type TEXT NOT NULL,
                sha256 TEXT NOT NULL UNIQUE,
                category TEXT NOT NULL,
                language TEXT NOT NULL,
                status TEXT NOT NULL DEFAULT 'active',
                file_path TEXT NOT NULL,
                extracted_text TEXT NOT NULL DEFAULT '',
                size_bytes INTEGER NOT NULL,
                created_at INTEGER NOT NULL,
                updated_at INTEGER NOT NULL
            );
            CREATE TABLE IF NOT EXISTS empleo_profile_facts (
                id TEXT PRIMARY KEY,
                category TEXT NOT NULL,
                fact_key TEXT NOT NULL,
                value_json TEXT NOT NULL,
                source_document_id TEXT NOT NULL,
                verified INTEGER NOT NULL DEFAULT 1,
                active INTEGER NOT NULL DEFAULT 1,
                created_at INTEGER NOT NULL,
                updated_at INTEGER NOT NULL,
                FOREIGN KEY(source_document_id) REFERENCES empleo_documents(id)
            );
            CREATE TABLE IF NOT EXISTS empleo_vacancies (
                id TEXT PRIMARY KEY,
                source_url TEXT NOT NULL DEFAULT '',
                status TEXT NOT NULL DEFAULT 'guardada',
                data_json TEXT NOT NULL,
                created_at INTEGER NOT NULL,
                updated_at INTEGER NOT NULL
            );
            CREATE TABLE IF NOT EXISTS empleo_generations (
                id TEXT PRIMARY KEY,
                vacancy_id TEXT,
                profile_track TEXT NOT NULL,
                language TEXT NOT NULL,
                content_json TEXT NOT NULL,
                status TEXT NOT NULL DEFAULT 'draft',
                artifacts_json TEXT NOT NULL DEFAULT '{}',
                created_at INTEGER NOT NULL,
                updated_at INTEGER NOT NULL,
                FOREIGN KEY(vacancy_id) REFERENCES empleo_vacancies(id)
            );
            CREATE INDEX IF NOT EXISTS idx_empleo_facts_source ON empleo_profile_facts(source_document_id);
            CREATE INDEX IF NOT EXISTS idx_empleo_vacancies_status ON empleo_vacancies(status);
            """
        )
    migrate_empleo_state()


def migrate_empleo_state() -> None:
    with db() as conn:
        if conn.execute("SELECT COUNT(*) FROM empleo_vacancies").fetchone()[0]:
            return
        row = conn.execute("SELECT state_json FROM app_state WHERE app_id='empleo'").fetchone()
        if not row:
            return
        try:
            vacancies = json.loads(row["state_json"]).get("vacantes", [])
        except (ValueError, TypeError):
            return
        for vacancy in vacancies:
            vacancy_id = str(vacancy.get("id") or uuid.uuid4())
            status = normalize_status(vacancy.get("status"))
            vacancy["id"] = vacancy_id
            vacancy["status"] = status
            conn.execute(
                "INSERT OR IGNORE INTO empleo_vacancies(id,source_url,status,data_json,created_at,updated_at) VALUES(?,?,?,?,?,?)",
                (
                    vacancy_id,
                    str(vacancy.get("link") or vacancy.get("source_url") or ""),
                    status,
                    json.dumps(vacancy, ensure_ascii=False),
                    int(vacancy.get("createdAt") or now()),
                    now(),
                ),
            )


def sign_session(username: str, issued_at: Optional[int] = None) -> str:
    ts = issued_at or now()
    secret = SESSION_SECRET or "session-secret-not-configured"
    payload = f"{username}:{ts}"
    signature = hmac.new(secret.encode(), payload.encode(), hashlib.sha256).hexdigest()
    return f"{payload}:{signature}"


def verify_session(token: str) -> bool:
    if not SESSION_SECRET:
        return False
    try:
        username, ts_text, signature = token.split(":", 2)
        ts = int(ts_text)
    except (ValueError, AttributeError):
        return False
    if username != WEB_USER or now() - ts > 60 * 60 * 24 * 14:
        return False
    expected = sign_session(username, ts).rsplit(":", 1)[1]
    return hmac.compare_digest(expected, signature)


def require_auth(
    x_app_token: Optional[str] = Header(None),
    x_session_token: Optional[str] = Header(None),
    session: Optional[str] = Cookie(None, alias=SESSION_COOKIE),
) -> None:
    if not AUTH_ENABLED:
        return
    if APP_TOKEN and x_app_token and hmac.compare_digest(APP_TOKEN, x_app_token):
        return
    if x_session_token and verify_session(x_session_token):
        return
    if session and verify_session(session):
        return
    raise HTTPException(status_code=401, detail="No autorizado")


def row_dict(row: sqlite3.Row) -> Dict[str, Any]:
    return {key: row[key] for key in row.keys()}


def normalize_status(value: Any) -> str:
    aliases = {
        "por_revisar": "guardada",
        "prioritaria": "prioritaria",
        "postulada": "postulada",
        "entrevista": "entrevista",
        "cerrada": "resultado",
        "cv": "cv_preparacion",
    }
    status = aliases.get(str(value), str(value or "guardada"))
    allowed = {"guardada", "prioritaria", "cv_preparacion", "postulada", "entrevista", "resultado"}
    return status if status in allowed else "guardada"


async def call_openrouter(
    prompt: str,
    system_prompt: str,
    is_json: bool = False,
    model: Optional[str] = None,
    max_tokens: int = 2400,
) -> Tuple[str, str]:
    if not OPENROUTER_API_KEY:
        raise HTTPException(status_code=500, detail="OPENROUTER_API_KEY no configurada")
    if is_json:
        prompt += "\n\nDevuelve unicamente JSON valido, sin markdown ni explicaciones."
    request_body = {
        "model": model or OPENROUTER_MODEL,
        "messages": [{"role": "system", "content": system_prompt}, {"role": "user", "content": prompt}],
        "temperature": 0.25,
        "max_tokens": max_tokens,
    }
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://juanpablogc.com",
        "X-Title": "Juan Pablo Agent Hub",
    }
    async with httpx.AsyncClient(timeout=90) as client:
        response = await client.post("https://openrouter.ai/api/v1/chat/completions", headers=headers, json=request_body)
    if response.status_code >= 400:
        raise HTTPException(status_code=response.status_code, detail=response.text[:500])
    data = response.json()
    return data["choices"][0]["message"]["content"], data.get("model") or request_body["model"]


def parse_ai_json(text: str) -> Dict[str, Any]:
    cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", text.strip(), flags=re.I)
    start, end = cleaned.find("{"), cleaned.rfind("}")
    if start < 0 or end <= start:
        raise HTTPException(status_code=502, detail="La IA no devolvio JSON valido")
    try:
        value = json.loads(cleaned[start : end + 1])
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=502, detail=f"JSON de IA invalido: {exc}")
    if not isinstance(value, dict):
        raise HTTPException(status_code=502, detail="La IA no devolvio un objeto JSON")
    return value


def classify_document(name: str) -> Tuple[str, str]:
    lowered = name.lower()
    language = "en" if re.search(r"(?:_|\b)en(?:_|\b|\.)", lowered) else "es"
    if "network" in lowered:
        return "networking", language
    if "roadmap" in lowered or "search_plan" in lowered or "job_search" in lowered:
        return "plan_busqueda", language
    if "linkedin" in lowered or "linkin" in lowered:
        return "linkedin", language
    if "cv" in lowered:
        if re.search(r"(^|[\\/ _-])da([\\/ _-]|$)", lowered):
            return "cv_data_analyst", language
        if re.search(r"(^|[\\/ _-])ds([\\/ _-]|$)", lowered):
            return "cv_data_scientist", language
        if re.search(r"(^|[\\/ _-])ml([\\/ _-]|$)", lowered):
            return "cv_machine_learning", language
        return "cv_general", language
    return "referencia", language


def extract_document_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        document = Document(str(path))
        blocks: List[str] = []
        blocks.extend(p.text.strip() for p in document.paragraphs if p.text.strip())
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if values:
                    blocks.append(" | ".join(values))
        return "\n\n".join(blocks)
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        return "\n\n".join((page.extract_text() or "").strip() for page in reader.pages).strip()
    return ""


def fact_category(text: str, fallback: str) -> str:
    lowered = text.lower()
    checks = [
        ("experiencia", ("experiencia", "instituto tecnologico", "space karani")),
        ("proyecto", ("proyecto", "churn", "clasificacion", "forecast")),
        ("habilidad", ("habilidad", "python", "sql", "machine learning")),
        ("educacion", ("educacion", "maestria", "ingenieria", "tripleten")),
        ("contacto", ("linkedin", "github", "gmail.com", "+52")),
    ]
    for category, needles in checks:
        if any(needle in lowered for needle in needles):
            return category
    return fallback


def process_document(document_id: str) -> Dict[str, Any]:
    with db() as conn:
        row = conn.execute("SELECT * FROM empleo_documents WHERE id=?", (document_id,)).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Documento no encontrado")
    path = Path(row["file_path"])
    if not path.is_file() or DOCUMENTS_DIR.resolve() not in path.resolve().parents:
        raise HTTPException(status_code=404, detail="Archivo privado no encontrado")
    try:
        extracted = extract_document_text(path)
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"No se pudo extraer el documento: {exc}")
    blocks = [block.strip() for block in re.split(r"\n\s*\n", extracted) if block.strip()]
    if not blocks and row["extension"] in {".png", ".jpg", ".jpeg"}:
        blocks = [f"Referencia visual privada: {row['original_name']}"]
    timestamp = now()
    with db() as conn:
        conn.execute("DELETE FROM empleo_profile_facts WHERE source_document_id=?", (document_id,))
        for index, block in enumerate(blocks[:160]):
            value = {"text": block[:3000], "document": row["original_name"]}
            conn.execute(
                "INSERT INTO empleo_profile_facts(id,category,fact_key,value_json,source_document_id,verified,active,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?)",
                (
                    str(uuid.uuid4()),
                    fact_category(block, row["category"]),
                    f"{row['category']}:{index + 1}",
                    json.dumps(value, ensure_ascii=False),
                    document_id,
                    1,
                    1,
                    timestamp,
                    timestamp,
                ),
            )
        conn.execute(
            "UPDATE empleo_documents SET extracted_text=?,updated_at=? WHERE id=?",
            (extracted[:250000], timestamp, document_id),
        )
    return {"id": document_id, "extracted_chars": len(extracted), "facts": len(blocks[:160])}


def import_document_path(source: Path, relative_name: Optional[str] = None) -> Dict[str, Any]:
    source = source.resolve()
    if not source.is_file() or source.suffix.lower() not in ALLOWED_DOCUMENT_EXTENSIONS:
        raise ValueError(f"Tipo de archivo no permitido: {source.name}")
    size = source.stat().st_size
    if size > MAX_UPLOAD_BYTES:
        raise ValueError(f"Archivo mayor a 10 MB: {source.name}")
    digest = hashlib.sha256(source.read_bytes()).hexdigest()
    with db() as conn:
        duplicate = conn.execute("SELECT * FROM empleo_documents WHERE sha256=?", (digest,)).fetchone()
    if duplicate:
        return {"document": row_dict(duplicate), "duplicate": True}
    document_id = str(uuid.uuid4())
    extension = source.suffix.lower()
    stored_name = f"{document_id}{extension}"
    destination = DOCUMENTS_DIR / stored_name
    shutil.copy2(str(source), str(destination))
    display_name = relative_name or source.name
    category, language = classify_document(display_name)
    timestamp = now()
    with db() as conn:
        conn.execute(
            "INSERT INTO empleo_documents(id,original_name,stored_name,extension,mime_type,sha256,category,language,status,file_path,size_bytes,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)",
            (
                document_id,
                display_name,
                stored_name,
                extension,
                ALLOWED_DOCUMENT_MIME[extension].__iter__().__next__(),
                digest,
                category,
                language,
                "active",
                str(destination),
                size,
                timestamp,
                timestamp,
            ),
        )
    processed = process_document(document_id)
    return {"document": get_document(document_id), "duplicate": False, "processed": processed}


def get_document(document_id: str) -> Dict[str, Any]:
    with db() as conn:
        row = conn.execute(
            "SELECT id,original_name,extension,mime_type,sha256,category,language,status,size_bytes,created_at,updated_at,LENGTH(extracted_text) extracted_chars FROM empleo_documents WHERE id=?",
            (document_id,),
        ).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Documento no encontrado")
    return row_dict(row)


def validate_public_url(raw_url: str) -> str:
    parsed = urlparse(raw_url.strip())
    if parsed.scheme not in {"http", "https"} or not parsed.hostname or parsed.username or parsed.password:
        raise HTTPException(status_code=422, detail="URL publica invalida")
    try:
        addresses = socket.getaddrinfo(parsed.hostname, parsed.port or (443 if parsed.scheme == "https" else 80))
    except socket.gaierror:
        raise HTTPException(status_code=422, detail="No se pudo resolver la URL")
    for address in addresses:
        ip = ipaddress.ip_address(address[4][0])
        if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_reserved or ip.is_multicast or ip.is_unspecified:
            raise HTTPException(status_code=422, detail="La URL apunta a una red no permitida")
    return parsed.geturl()


async def fetch_public_text(raw_url: str) -> Tuple[str, str]:
    current = validate_public_url(raw_url)
    headers = {"User-Agent": "Mozilla/5.0 (compatible; JuanPabloAgentHub/1.0)"}
    async with httpx.AsyncClient(timeout=12, follow_redirects=False, headers=headers) as client:
        for _ in range(4):
            async with client.stream("GET", current) as response:
                if response.status_code in {301, 302, 303, 307, 308}:
                    location = response.headers.get("location")
                    if not location:
                        raise HTTPException(status_code=422, detail="Redireccion sin destino")
                    current = validate_public_url(urljoin(current, location))
                    continue
                if response.status_code >= 400:
                    raise HTTPException(status_code=422, detail=f"El portal bloqueo la lectura ({response.status_code})")
                content_type = response.headers.get("content-type", "")
                if "text/html" not in content_type and "text/plain" not in content_type:
                    raise HTTPException(status_code=422, detail="El enlace no contiene texto web")
                chunks, total = [], 0
                async for chunk in response.aiter_bytes():
                    total += len(chunk)
                    if total > MAX_URL_BYTES:
                        raise HTTPException(status_code=422, detail="La pagina excede el limite permitido")
                    chunks.append(chunk)
                source = b"".join(chunks).decode(response.encoding or "utf-8", errors="replace")
                soup = BeautifulSoup(source, "html.parser")
                for tag in soup(["script", "style", "noscript", "svg"]):
                    tag.decompose()
                text = re.sub(r"\s+", " ", soup.get_text(" ", strip=True))
                return text[:80000], current
    raise HTTPException(status_code=422, detail="Demasiadas redirecciones")


def fallback_vacancy(text: str, url: str, request: VacancyParseRequest) -> Dict[str, Any]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return {
        "title": (lines[0] if lines else "Vacante por revisar")[:160],
        "company": "Por identificar",
        "location": request.location,
        "modality": request.modality,
        "salary": "No publicado",
        "contract_type": "No especificado",
        "language": "es" if request.language == "auto" else request.language,
        "level": request.level,
        "required_skills": [],
        "preferred_skills": [],
        "technologies": [],
        "keywords": [],
        "responsibilities": [],
        "match_score": 0,
        "gaps": ["Analisis con IA no disponible"],
        "next_action": "Revisar la ficha manualmente",
        "contact": "",
        "summary": text[:700],
        "source_url": url,
    }


def profile_source_text(track: str) -> Tuple[str, List[str]]:
    track_categories = {
        "general": {"cv_general", "referencia"},
        "da": {"cv_data_analyst", "cv_general"},
        "ds": {"cv_data_scientist", "cv_general"},
        "ml": {"cv_machine_learning", "cv_general"},
    }
    allowed = track_categories.get(track, set().union(*track_categories.values()))
    with db() as conn:
        rows = conn.execute(
            """
            SELECT f.id,f.value_json,d.category,d.original_name
            FROM empleo_profile_facts f
            JOIN empleo_documents d ON d.id=f.source_document_id
            WHERE f.verified=1 AND f.active=1 AND d.status='active'
              AND (d.category NOT LIKE 'cv_%' OR d.category IN ({}))
            ORDER BY CASE WHEN d.category IN ({}) THEN 0 ELSE 1 END,d.created_at,f.created_at
            LIMIT 240
            """.format(
                ",".join("?" for _ in allowed) or "''",
                ",".join("?" for _ in allowed) or "''",
            ),
            tuple(allowed) + tuple(allowed),
        ).fetchall()
        profile_row = conn.execute("SELECT state_json FROM app_state WHERE app_id='empleo_profile'").fetchone()
    pieces, ids, used = [], [], 0
    if profile_row:
        pieces.append("PERFIL EDITADO:\n" + profile_row["state_json"][:10000])
    for row in rows:
        value = json.loads(row["value_json"]).get("text", "")
        if not value or used + len(value) > 30000:
            continue
        pieces.append(f"[FUENTE {row['id']}] {row['original_name']}: {value}")
        ids.append(row["id"])
        used += len(value)
    return "\n\n".join(pieces), ids


def normalize_cv_content(content: Dict[str, Any]) -> Dict[str, Any]:
    def text_value(value: Any, limit: int = 1200) -> str:
        return str(value or "").strip()[:limit]

    def list_text(value: Any, count: int, limit: int = 500) -> List[str]:
        if not isinstance(value, list):
            return []
        return [text_value(item, limit) for item in value[:count] if text_value(item, limit)]

    normalized: Dict[str, Any] = {
        "name": text_value(content.get("name") or "Juan Pablo Garcia Chavez", 120),
        "headline": text_value(content.get("headline"), 180),
        "location": text_value(content.get("location") or "Morelia, Mexico", 100),
        "email": text_value(content.get("email"), 180),
        "phone": text_value(content.get("phone"), 80),
        "linkedin": text_value(content.get("linkedin"), 240),
        "github": text_value(content.get("github"), 240),
        "portfolio": text_value(content.get("portfolio"), 240),
        "summary": text_value(content.get("summary"), 1000),
        "skills": [],
        "experience": [],
        "projects": [],
        "education": list_text(content.get("education"), 5, 500),
        "languages": list_text(content.get("languages"), 4, 200),
        "cover_letter": text_value(content.get("cover_letter"), 4200),
        "keywords_used": list_text(content.get("keywords_used"), 30, 100),
        "keywords_missing": list_text(content.get("keywords_missing"), 20, 100),
        "warnings": list_text(content.get("warnings"), 20, 300),
        "source_fact_ids": list_text(content.get("source_fact_ids"), 80, 80),
    }
    for group in (content.get("skills") if isinstance(content.get("skills"), list) else [])[:6]:
        if isinstance(group, dict):
            normalized["skills"].append({"title": text_value(group.get("title"), 100), "items": list_text(group.get("items"), 18, 100)})
    for item in (content.get("experience") if isinstance(content.get("experience"), list) else [])[:5]:
        if isinstance(item, dict):
            normalized["experience"].append(
                {
                    "role": text_value(item.get("role"), 160),
                    "organization": text_value(item.get("organization"), 180),
                    "dates": text_value(item.get("dates"), 100),
                    "bullets": list_text(item.get("bullets"), 4, 380),
                }
            )
    for item in (content.get("projects") if isinstance(content.get("projects"), list) else [])[:5]:
        if isinstance(item, dict):
            normalized["projects"].append(
                {"name": text_value(item.get("name"), 180), "bullets": list_text(item.get("bullets"), 3, 380), "technologies": list_text(item.get("technologies"), 12, 100)}
            )
    return normalized


def safe_filename(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_-]+", "_", value).strip("_")
    return cleaned[:70] or "CV"


def add_docx_hyperlink(paragraph: Any, label: str, url: str) -> None:
    if not url:
        return
    relationship_id = paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(color)
    properties.append(underline)
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = label
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_docx(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.52)
    section.bottom_margin = Inches(0.52)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.0
    for style_name, size in (("Title", 18), ("Heading 1", 12), ("Heading 2", 10)):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)


def add_docx_section(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)


def build_cv_docx(content: Dict[str, Any], path: Path) -> None:
    document = Document()
    configure_docx(document)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(1)
    run = title.add_run(f"{content['name']} | {content['headline']}")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(1)
    contact.add_run(" | ".join(item for item in [content["location"], content["email"], content["phone"]] if item))
    links = document.add_paragraph()
    links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for index, (label, url) in enumerate((("LinkedIn", content["linkedin"]), ("GitHub", content["github"]), ("Portafolio", content["portfolio"]))):
        if url:
            if index:
                links.add_run(" | ")
            add_docx_hyperlink(links, label, url)
    summary = document.add_paragraph(content["summary"])
    summary.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if content["skills"]:
        add_docx_section(document, "Habilidades tecnicas")
        for group in content["skills"]:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(group["title"] + ": ").bold = True
            paragraph.add_run(" | ".join(group["items"]))
    if content["experience"]:
        add_docx_section(document, "Experiencia profesional")
        for item in content["experience"]:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.keep_with_next = True
            paragraph.add_run(item["role"]).bold = True
            if item["organization"]:
                paragraph.add_run("\n" + item["organization"]).bold = True
            if item["dates"]:
                paragraph.add_run(" | " + item["dates"]).bold = True
            for bullet in item["bullets"]:
                p = document.add_paragraph(bullet, style="List Bullet")
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.first_line_indent = Inches(-0.14)
    if content["projects"]:
        add_docx_section(document, "Proyectos relevantes")
        for item in content["projects"]:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.keep_with_next = True
            paragraph.add_run(item["name"]).bold = True
            for bullet in item["bullets"]:
                p = document.add_paragraph(bullet, style="List Bullet")
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.first_line_indent = Inches(-0.14)
            if item["technologies"]:
                p = document.add_paragraph()
                p.add_run("Tecnologias: ").bold = True
                p.add_run(", ".join(item["technologies"]))
    if content["education"]:
        add_docx_section(document, "Educacion")
        for value in content["education"]:
            document.add_paragraph(value)
    if content["languages"]:
        add_docx_section(document, "Idiomas")
        document.add_paragraph(" | ".join(content["languages"])).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.save(str(path))


def build_letter_docx(content: Dict[str, Any], vacancy: Dict[str, Any], path: Path) -> None:
    document = Document()
    configure_docx(document)
    title = document.add_paragraph()
    title.add_run(content["name"]).bold = True
    document.add_paragraph(" | ".join(item for item in [content["email"], content["phone"], content["linkedin"]] if item))
    document.add_paragraph(time.strftime("%d/%m/%Y"))
    document.add_paragraph(f"Equipo de seleccion\n{vacancy.get('company') or 'Empresa objetivo'}")
    document.add_paragraph(f"Asunto: Postulacion a {vacancy.get('title') or content['headline']}")
    for block in cover_letter_blocks(content["cover_letter"], content["name"]):
        if block:
            paragraph = document.add_paragraph(block)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    document.add_paragraph("Atentamente,\n" + content["name"])
    document.save(str(path))


def pdf_styles() -> Dict[str, ParagraphStyle]:
    return {
        "title": ParagraphStyle("CVTitle", fontName="Helvetica-Bold", fontSize=16.5, leading=18, alignment=TA_CENTER, spaceAfter=2),
        "contact": ParagraphStyle("Contact", fontName="Helvetica", fontSize=8.5, leading=10, alignment=TA_CENTER, spaceAfter=4),
        "summary": ParagraphStyle("Summary", fontName="Helvetica", fontSize=8.7, leading=10.5, alignment=TA_JUSTIFY, spaceAfter=3),
        "section": ParagraphStyle("Section", fontName="Helvetica-Bold", fontSize=10.5, leading=12, alignment=TA_CENTER, spaceBefore=4, spaceAfter=2),
        "heading": ParagraphStyle("Heading", fontName="Helvetica-Bold", fontSize=8.8, leading=10.2, alignment=TA_LEFT, spaceBefore=2, spaceAfter=0),
        "body": ParagraphStyle("Body", fontName="Helvetica", fontSize=8.3, leading=9.8, alignment=TA_LEFT, spaceAfter=1.5),
        "bullet": ParagraphStyle("Bullet", fontName="Helvetica", fontSize=8.2, leading=9.6, leftIndent=13, firstLineIndent=-7, bulletIndent=5, spaceAfter=1),
    }


def escape_pdf(value: Any) -> str:
    return html.escape(str(value or ""), quote=True)


def build_cv_pdf(content: Dict[str, Any], path: Path) -> int:
    styles = pdf_styles()
    story: List[Any] = []
    story.append(Paragraph(escape_pdf(content["name"] + " | " + content["headline"]), styles["title"]))
    contact = " | ".join(escape_pdf(item) for item in [content["location"], content["email"], content["phone"]] if item)
    story.append(Paragraph(contact, styles["contact"]))
    link_parts = []
    for label, url in (("LinkedIn", content["linkedin"]), ("GitHub", content["github"]), ("Portafolio", content["portfolio"])):
        if url:
            link_parts.append(f'<link href="{escape_pdf(url)}" color="#0563C1">{label}</link>')
    if link_parts:
        story.append(Paragraph(" | ".join(link_parts), styles["contact"]))
    story.append(Paragraph(escape_pdf(content["summary"]), styles["summary"]))
    if content["skills"]:
        story.append(Paragraph("Habilidades tecnicas", styles["section"]))
        for group in content["skills"]:
            story.append(Paragraph(f"<b>{escape_pdf(group['title'])}:</b> {escape_pdf(' | '.join(group['items']))}", styles["body"]))
    if content["experience"]:
        story.append(Paragraph("Experiencia profesional", styles["section"]))
        for item in content["experience"]:
            heading = f"<b>{escape_pdf(item['role'])}</b>"
            if item["organization"]:
                heading += f"<br/><b>{escape_pdf(item['organization'])}</b>"
            if item["dates"]:
                heading += f" | <b>{escape_pdf(item['dates'])}</b>"
            story.append(Paragraph(heading, styles["heading"]))
            for bullet in item["bullets"]:
                story.append(Paragraph(escape_pdf(bullet), styles["bullet"], bulletText="•"))
    if content["projects"]:
        story.append(Paragraph("Proyectos relevantes", styles["section"]))
        for item in content["projects"]:
            story.append(Paragraph(f"<b>{escape_pdf(item['name'])}</b>", styles["heading"]))
            for bullet in item["bullets"]:
                story.append(Paragraph(escape_pdf(bullet), styles["bullet"], bulletText="•"))
            if item["technologies"]:
                story.append(Paragraph(f"<b>Tecnologias:</b> {escape_pdf(', '.join(item['technologies']))}", styles["body"]))
    if content["education"]:
        story.append(Paragraph("Educacion", styles["section"]))
        for value in content["education"]:
            story.append(Paragraph(escape_pdf(value), styles["body"]))
    if content["languages"]:
        story.append(Paragraph("Idiomas", styles["section"]))
        story.append(Paragraph(escape_pdf(" | ".join(content["languages"])), styles["contact"]))
    document = SimpleDocTemplate(str(path), pagesize=letter, rightMargin=0.52 * inch, leftMargin=0.52 * inch, topMargin=0.45 * inch, bottomMargin=0.45 * inch, title=f"CV {content['name']}")
    document.build(story)
    pages = len(PdfReader(str(path)).pages)
    if pages > 2:
        path.unlink(missing_ok=True)
        raise HTTPException(status_code=422, detail="El CV excede dos paginas; reduce el borrador antes de aprobar")
    return pages


def build_letter_pdf(content: Dict[str, Any], vacancy: Dict[str, Any], path: Path) -> int:
    styles = getSampleStyleSheet()
    body = ParagraphStyle("LetterBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5, leading=14, alignment=TA_JUSTIFY, spaceAfter=10)
    contact_style = ParagraphStyle("LetterContact", parent=body, alignment=TA_LEFT)
    heading = ParagraphStyle("LetterHeading", parent=body, fontName="Helvetica-Bold", alignment=TA_LEFT)
    story: List[Any] = [
        Paragraph(escape_pdf(content["name"]), ParagraphStyle("Name", parent=heading, fontSize=15, leading=18)),
        Paragraph(escape_pdf(" | ".join(item for item in [content["email"], content["phone"], content["linkedin"]] if item)), contact_style),
        Spacer(1, 12),
        Paragraph(time.strftime("%d/%m/%Y"), body),
        Paragraph(f"Equipo de seleccion<br/>{escape_pdf(vacancy.get('company') or 'Empresa objetivo')}", body),
        Paragraph(f"Asunto: Postulacion a {escape_pdf(vacancy.get('title') or content['headline'])}", heading),
    ]
    for block in cover_letter_blocks(content["cover_letter"], content["name"]):
        if block:
            story.append(Paragraph(escape_pdf(block), body))
    story.append(Paragraph("Atentamente,<br/>" + escape_pdf(content["name"]), body))
    document = SimpleDocTemplate(str(path), pagesize=letter, rightMargin=0.75 * inch, leftMargin=0.75 * inch, topMargin=0.7 * inch, bottomMargin=0.7 * inch, title="Carta de presentacion")
    document.build(story)
    pages = len(PdfReader(str(path)).pages)
    if pages > 2:
        path.unlink(missing_ok=True)
        raise HTTPException(status_code=422, detail="La carta excede dos paginas; reduce el borrador")
    return pages


def cover_letter_blocks(value: str, name: str) -> List[str]:
    blocks = [block.strip() for block in re.split(r"\n\s*\n", value or "") if block.strip()]
    while blocks and (
        blocks[-1].lower().rstrip(".,") in {"atentamente", "saludos", "saludos cordiales", "sincerely", "best regards"}
        or blocks[-1].lower().strip() == name.lower().strip()
    ):
        blocks.pop()
    return blocks


def create_generation_artifacts(generation_id: str, content: Dict[str, Any], vacancy: Dict[str, Any]) -> Dict[str, Any]:
    folder = (ARTIFACTS_DIR / generation_id).resolve()
    if ARTIFACTS_DIR.resolve() not in folder.parents:
        raise HTTPException(status_code=500, detail="Ruta de salida invalida")
    folder.mkdir(parents=True, exist_ok=True)
    stem = safe_filename(f"CV_{content['name']}_{vacancy.get('company') or 'Empresa'}")
    files = {
        "cv_pdf": (str(uuid.uuid4()), folder / f"{stem}.pdf"),
        "cv_docx": (str(uuid.uuid4()), folder / f"{stem}.docx"),
        "letter_pdf": (str(uuid.uuid4()), folder / f"Carta_{stem}.pdf"),
        "letter_docx": (str(uuid.uuid4()), folder / f"Carta_{stem}.docx"),
    }
    build_cv_docx(content, files["cv_docx"][1])
    cv_pages = build_cv_pdf(content, files["cv_pdf"][1])
    build_letter_docx(content, vacancy, files["letter_docx"][1])
    letter_pages = build_letter_pdf(content, vacancy, files["letter_pdf"][1])
    artifacts: Dict[str, Any] = {}
    for key, (file_id, path) in files.items():
        artifacts[key] = {"id": file_id, "name": path.name, "path": str(path), "size": path.stat().st_size, "url": f"/api/empleo/files/{file_id}"}
    artifacts["cv_pdf"]["pages"] = cv_pages
    artifacts["letter_pdf"]["pages"] = letter_pages
    return artifacts


@app.on_event("startup")
def startup() -> None:
    init_db()


@app.get("/api/health")
def health() -> Dict[str, Any]:
    return {
        "status": "ok",
        "version": app.version,
        "openrouter_configured": bool(OPENROUTER_API_KEY),
        "model": OPENROUTER_MODEL,
        "database": str(DB_PATH),
        "employment_documents": DOCUMENTS_DIR.exists(),
    }


@app.post("/api/auth/login")
def login(payload: LoginRequest, response: Response) -> Dict[str, Any]:
    if not WEB_PASSWORD or not SESSION_SECRET:
        raise HTTPException(status_code=503, detail="Autenticacion no configurada")
    if payload.username != WEB_USER or not hmac.compare_digest(payload.password, WEB_PASSWORD):
        raise HTTPException(status_code=401, detail="Credenciales incorrectas")
    session_token = sign_session(WEB_USER)
    response.set_cookie(SESSION_COOKIE, session_token, httponly=True, samesite="lax", secure=COOKIE_SECURE, max_age=60 * 60 * 24 * 14)
    return {"ok": True, "session_token": session_token}


@app.post("/api/auth/logout")
def logout(response: Response) -> Dict[str, Any]:
    response.delete_cookie(SESSION_COOKIE)
    return {"ok": True}


@app.get("/api/auth/status")
def auth_status(x_session_token: Optional[str] = Header(None), session: Optional[str] = Cookie(None, alias=SESSION_COOKIE)) -> Dict[str, Any]:
    return {"authenticated": bool((x_session_token and verify_session(x_session_token)) or (session and verify_session(session)))}


@app.post("/api/ai/chat")
async def ai_chat(payload: AIRequest, _: None = Depends(require_auth)) -> Dict[str, Any]:
    text, model = await call_openrouter(
        payload.prompt,
        payload.system_prompt or "Eres un asistente experto, claro y practico. Responde en espanol.",
        payload.is_json,
        payload.model,
        1800,
    )
    return {"text": text, "model": model}


@app.post("/api/orchestrator/chat")
async def orchestrator_chat(payload: OrchestratorRequest, _: None = Depends(require_auth)) -> Dict[str, Any]:
    import asyncio
    import re
    
    cmd = ["/usr/local/bin/picoclaw", "agent", "-m", payload.prompt]
    if payload.session_id:
        cmd.extend(["-s", payload.session_id])
        
    try:
        process = await asyncio.create_subprocess_exec(
            *cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE
        )
        stdout, stderr = await process.communicate()
        
        if process.returncode != 0:
            return {
                "text": f"Error de ejecución en PicoClaw: {stderr.decode('utf-8', errors='ignore').strip()}",
                "model": "picoclaw-error"
            }
            
        raw_text = stdout.decode('utf-8', errors='ignore').strip()
        
        if "🦞" in raw_text:
            base_text = raw_text.split("🦞")[-1].strip()
        else:
            base_text = raw_text
            
        # Filtrar líneas de logs que contengan marcas de tiempo de Go
        clean_lines = [
            line for line in base_text.splitlines()
            if not re.match(r'^\d{2}:\d{2}:\d{2}\s+(INF|DBG|WRN|ERR)', line.strip())
        ]
        clean_text = "\n".join(clean_lines).strip()
        
        return {"text": clean_text, "model": "picoclaw-subprocess"}
        
    except Exception as e:
        return {"text": f"Excepción en el servidor FastAPI: {str(e)}", "model": "fastapi-exception"}


@app.get("/api/state/{app_id}")
def get_state(app_id: str, _: None = Depends(require_auth)) -> Dict[str, Any]:
    with db() as conn:
        row = conn.execute("SELECT state_json,updated_at FROM app_state WHERE app_id=?", (app_id,)).fetchone()
    return {"state": json.loads(row["state_json"]), "updated_at": row["updated_at"]} if row else {"state": None, "updated_at": None}


@app.put("/api/state/{app_id}")
def put_state(app_id: str, payload: StateRequest, _: None = Depends(require_auth)) -> Dict[str, Any]:
    updated_at = now()
    with db() as conn:
        conn.execute(
            "INSERT INTO app_state(app_id,state_json,updated_at) VALUES(?,?,?) ON CONFLICT(app_id) DO UPDATE SET state_json=excluded.state_json,updated_at=excluded.updated_at",
            (app_id, json.dumps(payload.state, ensure_ascii=False), updated_at),
        )
    return {"ok": True, "updated_at": updated_at}


@app.post("/api/bitacora")
def save_bitacora(payload: BitacoraRequest, _: None = Depends(require_auth)) -> Dict[str, Any]:
    created_at = now()
    with db() as conn:
        cursor = conn.execute("INSERT INTO bitacoras(app_id,fecha,contenido,created_at) VALUES(?,?,?,?)", (payload.app_id, payload.fecha, payload.contenido, created_at))
    return {"ok": True, "id": cursor.lastrowid, "created_at": created_at}


@app.get("/api/empleo/documents")
def list_documents(_: None = Depends(require_auth)) -> Dict[str, Any]:
    with db() as conn:
        rows = conn.execute(
            "SELECT id,original_name,extension,mime_type,sha256,category,language,status,size_bytes,created_at,updated_at,LENGTH(extracted_text) extracted_chars FROM empleo_documents ORDER BY created_at DESC"
        ).fetchall()
    return {"documents": [row_dict(row) for row in rows]}


@app.post("/api/empleo/documents")
async def upload_document(file: UploadFile = File(...), category: str = Form("auto"), _: None = Depends(require_auth)) -> Dict[str, Any]:
    original_name = Path(file.filename or "document").name
    extension = Path(original_name).suffix.lower()
    if extension not in ALLOWED_DOCUMENT_EXTENSIONS:
        raise HTTPException(status_code=415, detail="Solo se permiten DOCX, PDF, PNG y JPG")
    if file.content_type not in ALLOWED_DOCUMENT_MIME[extension]:
        raise HTTPException(status_code=415, detail="El tipo MIME no coincide con el archivo")
    document_id = str(uuid.uuid4())
    stored_name = f"{document_id}{extension}"
    destination = DOCUMENTS_DIR / stored_name
    digest = hashlib.sha256()
    size = 0
    try:
        with destination.open("wb") as output:
            while True:
                chunk = await file.read(1024 * 1024)
                if not chunk:
                    break
                size += len(chunk)
                if size > MAX_UPLOAD_BYTES:
                    raise HTTPException(status_code=413, detail="El archivo excede 10 MB")
                digest.update(chunk)
                output.write(chunk)
    except Exception:
        destination.unlink(missing_ok=True)
        raise
    finally:
        await file.close()
    sha256 = digest.hexdigest()
    with db() as conn:
        duplicate = conn.execute("SELECT id FROM empleo_documents WHERE sha256=?", (sha256,)).fetchone()
        if duplicate:
            destination.unlink(missing_ok=True)
            return {"document": get_document(duplicate["id"]), "duplicate": True}
    detected_category, language = classify_document(original_name)
    timestamp = now()
    with db() as conn:
        conn.execute(
            "INSERT INTO empleo_documents(id,original_name,stored_name,extension,mime_type,sha256,category,language,status,file_path,size_bytes,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)",
            (document_id, original_name, stored_name, extension, file.content_type or "application/octet-stream", sha256, detected_category if category == "auto" else category, language, "active", str(destination), size, timestamp, timestamp),
        )
    processed = process_document(document_id)
    return {"document": get_document(document_id), "duplicate": False, "processed": processed}


@app.get("/api/empleo/documents/{document_id}/download")
def download_document(document_id: str, _: None = Depends(require_auth)) -> FileResponse:
    with db() as conn:
        row = conn.execute("SELECT original_name,mime_type,file_path FROM empleo_documents WHERE id=?", (document_id,)).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Documento no encontrado")
    path = Path(row["file_path"]).resolve()
    if not path.is_file() or DOCUMENTS_DIR.resolve() not in path.parents:
        raise HTTPException(status_code=404, detail="Archivo no encontrado")
    return FileResponse(str(path), media_type=row["mime_type"], filename=Path(row["original_name"]).name)


@app.post("/api/empleo/documents/{document_id}/process")
def reprocess_document(document_id: str, _: None = Depends(require_auth)) -> Dict[str, Any]:
    return process_document(document_id)


@app.patch("/api/empleo/documents/{document_id}")
def patch_document(document_id: str, payload: DocumentPatch, _: None = Depends(require_auth)) -> Dict[str, Any]:
    if payload.status and payload.status not in {"active", "archived"}:
        raise HTTPException(status_code=422, detail="Estado de documento invalido")
    fields, values = [], []
    if payload.category is not None:
        fields.append("category=?")
        values.append(payload.category[:80])
    if payload.status is not None:
        fields.append("status=?")
        values.append(payload.status)
    if not fields:
        return get_document(document_id)
    fields.append("updated_at=?")
    values.extend([now(), document_id])
    with db() as conn:
        result = conn.execute(f"UPDATE empleo_documents SET {','.join(fields)} WHERE id=?", tuple(values))
        if not result.rowcount:
            raise HTTPException(status_code=404, detail="Documento no encontrado")
    return get_document(document_id)


@app.get("/api/empleo/profile")
def get_profile(_: None = Depends(require_auth)) -> Dict[str, Any]:
    with db() as conn:
        profile_row = conn.execute("SELECT state_json,updated_at FROM app_state WHERE app_id='empleo_profile'").fetchone()
        facts = conn.execute(
            "SELECT f.id,f.category,f.fact_key,f.value_json,f.verified,f.active,f.source_document_id,d.original_name FROM empleo_profile_facts f JOIN empleo_documents d ON d.id=f.source_document_id WHERE d.status='active' ORDER BY d.category,f.created_at"
        ).fetchall()
        document_count = conn.execute("SELECT COUNT(*) FROM empleo_documents WHERE status='active'").fetchone()[0]
    return {
        "profile": json.loads(profile_row["state_json"]) if profile_row else {},
        "updated_at": profile_row["updated_at"] if profile_row else None,
        "document_count": document_count,
        "facts": [{**row_dict(row), "value": json.loads(row["value_json"])} for row in facts],
    }


@app.put("/api/empleo/profile")
def put_profile(payload: ProfileRequest, _: None = Depends(require_auth)) -> Dict[str, Any]:
    updated_at = now()
    with db() as conn:
        conn.execute(
            "INSERT INTO app_state(app_id,state_json,updated_at) VALUES('empleo_profile',?,?) ON CONFLICT(app_id) DO UPDATE SET state_json=excluded.state_json,updated_at=excluded.updated_at",
            (json.dumps(payload.profile, ensure_ascii=False), updated_at),
        )
    return {"ok": True, "updated_at": updated_at}


@app.get("/api/empleo/vacancies")
def list_vacancies(_: None = Depends(require_auth)) -> Dict[str, Any]:
    with db() as conn:
        rows = conn.execute("SELECT id,status,data_json,created_at,updated_at FROM empleo_vacancies ORDER BY updated_at DESC").fetchall()
    vacancies = []
    for row in rows:
        value = json.loads(row["data_json"])
        value.update({"id": row["id"], "status": row["status"], "created_at": row["created_at"], "updated_at": row["updated_at"]})
        vacancies.append(value)
    return {"vacancies": vacancies}


@app.post("/api/empleo/vacancies/parse")
async def parse_vacancy(payload: VacancyParseRequest, _: None = Depends(require_auth)) -> Dict[str, Any]:
    raw_text = payload.text.strip()
    fetched_text, final_url, fetch_warning = "", payload.url.strip(), ""
    if payload.url.strip():
        try:
            fetched_text, final_url = await fetch_public_text(payload.url)
        except HTTPException as exc:
            fetch_warning = str(exc.detail)
    if not raw_text and not fetched_text:
        raise HTTPException(status_code=422, detail=fetch_warning or "Pega la vacante o proporciona un enlace publico")
    source = raw_text
    if fetched_text:
        source += ("\n\nCONTENIDO DEL ENLACE:\n" + fetched_text[:30000])
    prompt = f"""Extrae una vacante en una ficha profesional. El texto pegado por el usuario es la fuente autoritativa.
Perfil objetivo: {payload.profile_track}; modalidad: {payload.modality}; ubicacion: {payload.location}; nivel: {payload.level}.
Devuelve exactamente estas claves: title, company, location, modality, salary, contract_type, language, level, required_skills, preferred_skills, technologies, keywords, responsibilities, match_score, gaps, next_action, contact, summary.
Las listas deben ser arreglos de texto; match_score debe ser 0-100. No inventes salario, empresa ni requisitos ausentes.

VACANTE:
{source[:50000]}"""
    try:
        answer, model = await call_openrouter(prompt, "Eres un analista ATS riguroso. Extraes datos sin inventar informacion.", True, max_tokens=2200)
        vacancy = parse_ai_json(answer)
    except HTTPException as exc:
        if exc.status_code in {500, 502, 503}:
            vacancy = fallback_vacancy(raw_text or fetched_text, final_url, payload)
            model = None
            fetch_warning = (fetch_warning + "; " if fetch_warning else "") + str(exc.detail)
        else:
            raise
    vacancy_id = str(uuid.uuid4())
    vacancy.update(
        {
            "id": vacancy_id,
            "source_url": final_url,
            "raw_text": raw_text[:60000],
            "profile_track": payload.profile_track,
            "status": "guardada",
            "fetch_warning": fetch_warning,
        }
    )
    timestamp = now()
    with db() as conn:
        conn.execute(
            "INSERT INTO empleo_vacancies(id,source_url,status,data_json,created_at,updated_at) VALUES(?,?,?,?,?,?)",
            (vacancy_id, final_url, "guardada", json.dumps(vacancy, ensure_ascii=False), timestamp, timestamp),
        )
    return {"vacancy": vacancy, "model": model, "fetch_warning": fetch_warning}


@app.patch("/api/empleo/vacancies/{vacancy_id}")
def patch_vacancy(vacancy_id: str, payload: VacancyPatch, _: None = Depends(require_auth)) -> Dict[str, Any]:
    with db() as conn:
        row = conn.execute("SELECT data_json,status FROM empleo_vacancies WHERE id=?", (vacancy_id,)).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Vacante no encontrada")
        data = json.loads(row["data_json"])
        allowed = {"title", "company", "location", "modality", "salary", "contract_type", "language", "level", "required_skills", "preferred_skills", "technologies", "keywords", "responsibilities", "match_score", "gaps", "next_action", "contact", "summary", "notes", "status"}
        for key, value in payload.fields.items():
            if key in allowed:
                data[key] = value
        status = normalize_status(data.get("status") or row["status"])
        data["status"] = status
        conn.execute("UPDATE empleo_vacancies SET status=?,data_json=?,updated_at=? WHERE id=?", (status, json.dumps(data, ensure_ascii=False), now(), vacancy_id))
    return {"vacancy": data}


@app.post("/api/empleo/cv/generate")
async def generate_cv(payload: CVGenerateRequest, _: None = Depends(require_auth)) -> Dict[str, Any]:
    vacancy = payload.vacancy
    requested_track = payload.profile_track
    if requested_track == "auto":
        text = " ".join(str(vacancy.get(key, "")) for key in ("title", "summary", "technologies", "keywords")).lower()
        requested_track = "ml" if "machine learning" in text or "ml engineer" in text else "da" if "data analyst" in text or "business intelligence" in text else "ds" if "data scientist" in text else "general"
    language = payload.language
    if language == "auto":
        language = "en" if str(vacancy.get("language", "")).lower().startswith("en") else "es"
    source_text, allowed_fact_ids = profile_source_text(requested_track)
    if not source_text:
        raise HTTPException(status_code=422, detail="La biblioteca profesional esta vacia; importa y procesa los documentos primero")
    prompt = f"""Genera un borrador ATS para Juan Pablo usando SOLO los hechos de FUENTES autorizadas. No inventes empleos, fechas, estudios, tecnologias ni metricas. Adapta el vocabulario sin alterar hechos.
Idioma: {language}. Perfil: {requested_track}. CV de una columna, tamano Carta, maximo dos paginas. Maximo 5 experiencias con 4 bullets, 5 proyectos con 3 bullets y 6 grupos de habilidades. Redacta tambien una carta breve.
Devuelve estas claves JSON: name, headline, location, email, phone, linkedin, github, portfolio, summary, skills, experience, projects, education, languages, cover_letter, keywords_used, keywords_missing, warnings, source_fact_ids.
skills: [{{title,items}}]; experience: [{{role,organization,dates,bullets}}]; projects: [{{name,bullets,technologies}}]. source_fact_ids solo puede contener IDs escritos como FUENTE.

VACANTE:
{json.dumps(vacancy, ensure_ascii=False)[:22000]}

FUENTES AUTORIZADAS:
{source_text}"""
    answer, model = await call_openrouter(prompt, "Eres un redactor de CV ATS extremadamente riguroso con la veracidad. Si un dato no esta en las fuentes, lo omites y agregas una advertencia.", True, max_tokens=3900)
    content = normalize_cv_content(parse_ai_json(answer))
    valid_ids = set(allowed_fact_ids)
    invalid_ids = [fact_id for fact_id in content["source_fact_ids"] if fact_id not in valid_ids]
    content["source_fact_ids"] = [fact_id for fact_id in content["source_fact_ids"] if fact_id in valid_ids]
    if invalid_ids:
        content["warnings"].append("Se eliminaron referencias de fuente no validas.")
    generation_id = str(uuid.uuid4())
    timestamp = now()
    vacancy_id = str(vacancy.get("id") or "") or None
    with db() as conn:
        if vacancy_id and not conn.execute("SELECT 1 FROM empleo_vacancies WHERE id=?", (vacancy_id,)).fetchone():
            vacancy_id = None
        conn.execute(
            "INSERT INTO empleo_generations(id,vacancy_id,profile_track,language,content_json,status,artifacts_json,created_at,updated_at) VALUES(?,?,?,?,?,'draft','{}',?,?)",
            (generation_id, vacancy_id, requested_track, language, json.dumps(content, ensure_ascii=False), timestamp, timestamp),
        )
        if vacancy_id:
            conn.execute("UPDATE empleo_vacancies SET status='cv_preparacion',updated_at=? WHERE id=?", (timestamp, vacancy_id))
    return {"generation": {"id": generation_id, "status": "draft", "profile_track": requested_track, "language": language, "content": content}, "model": model}


@app.patch("/api/empleo/cv/{generation_id}")
def patch_generation(generation_id: str, payload: GenerationPatch, _: None = Depends(require_auth)) -> Dict[str, Any]:
    content = normalize_cv_content(payload.content)
    with db() as conn:
        result = conn.execute("UPDATE empleo_generations SET content_json=?,updated_at=? WHERE id=? AND status='draft'", (json.dumps(content, ensure_ascii=False), now(), generation_id))
        if not result.rowcount:
            raise HTTPException(status_code=404, detail="Borrador no encontrado o ya aprobado")
    return {"generation": {"id": generation_id, "status": "draft", "content": content}}


@app.post("/api/empleo/cv/{generation_id}/approve")
def approve_generation(generation_id: str, _: None = Depends(require_auth)) -> Dict[str, Any]:
    with db() as conn:
        row = conn.execute("SELECT * FROM empleo_generations WHERE id=?", (generation_id,)).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Generacion no encontrada")
        if row["status"] == "approved":
            return {"generation": {"id": generation_id, "status": "approved", "artifacts": json.loads(row["artifacts_json"])}}
        content = normalize_cv_content(json.loads(row["content_json"]))
        vacancy = {}
        if row["vacancy_id"]:
            vacancy_row = conn.execute("SELECT data_json FROM empleo_vacancies WHERE id=?", (row["vacancy_id"],)).fetchone()
            if vacancy_row:
                vacancy = json.loads(vacancy_row["data_json"])
    artifacts = create_generation_artifacts(generation_id, content, vacancy)
    with db() as conn:
        conn.execute("UPDATE empleo_generations SET status='approved',artifacts_json=?,updated_at=? WHERE id=?", (json.dumps(artifacts, ensure_ascii=False), now(), generation_id))
    public_artifacts = {key: {field: value for field, value in artifact.items() if field != "path"} for key, artifact in artifacts.items()}
    return {"generation": {"id": generation_id, "status": "approved", "artifacts": public_artifacts}}


@app.get("/api/empleo/files/{file_id}")
def download_generated_file(file_id: str, _: None = Depends(require_auth)) -> FileResponse:
    with db() as conn:
        rows = conn.execute("SELECT artifacts_json FROM empleo_generations WHERE status='approved'").fetchall()
    for row in rows:
        for artifact in json.loads(row["artifacts_json"]).values():
            if artifact.get("id") != file_id:
                continue
            path = Path(artifact["path"]).resolve()
            if not path.is_file() or ARTIFACTS_DIR.resolve() not in path.parents:
                raise HTTPException(status_code=404, detail="Archivo generado no encontrado")
            media_type = "application/pdf" if path.suffix.lower() == ".pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            return FileResponse(str(path), media_type=media_type, filename=artifact["name"])
    raise HTTPException(status_code=404, detail="Archivo generado no encontrado")


if not PUBLIC_DIR.exists():
    PUBLIC_DIR.mkdir(parents=True, exist_ok=True)

app.mount("/", StaticFiles(directory=str(PUBLIC_DIR), html=True), name="public")
